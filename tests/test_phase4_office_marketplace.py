"""Phase-4 tests: Office COM tools (fallbacks) + cloud plugin registry."""

from __future__ import annotations

import json
import sys
from pathlib import Path
from types import SimpleNamespace

import pytest

from james.core import workspace
from james.tools.office_tools import (
    _excel_openpyxl_read,
    _excel_openpyxl_write,
    excel_read_cells,
    excel_write_cells,
    word_read_document,
)
from james.tools.registry import ToolRegistry


@pytest.fixture
def ws(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> Path:
    monkeypatch.setattr(workspace.settings.assistant, "workspace_dir", tmp_path)
    return tmp_path


# ---------------------------------------------------------------------------
# excel fallbacks (no COM in CI -> openpyxl path)
# ---------------------------------------------------------------------------


def test_excel_openpyxl_write_then_read(ws: Path) -> None:
    path = ws / "book.xlsx"
    _excel_openpyxl_write(
        path, [["Name", "Score"], ["Ada", 9], ["Grace", 10]], sheet="Sheet1", start_cell="A1"
    )
    assert path.exists()
    rows = _excel_openpyxl_read(path, sheet="Sheet1", cell_range="A1:B3")
    assert rows == [["Name", "Score"], ["Ada", "9"], ["Grace", "10"]]


def test_excel_write_cells_tool_roundtrip(ws: Path) -> None:
    written = excel_write_cells.run(path="book.xlsx", data=json.dumps([[1, 2], [3, 4]]))
    assert written.ok
    read = excel_read_cells.run(path="book.xlsx", range="A1:B2")
    assert read.ok
    # COM Excel returns floats ("1.0"), openpyxl returns "1" — accept both.
    assert read.data == [["1", "2"], ["3", "4"]] or read.data == [["1.0", "2.0"], ["3.0", "4.0"]]


def test_excel_read_cells_missing_file(ws: Path) -> None:
    result = excel_read_cells.run(path="nope.xlsx")
    assert not result.ok and "not found" in result.output


def test_excel_write_invalid_data(ws: Path) -> None:
    result = excel_write_cells.run(path="b.xlsx", data="not json")
    assert not result.ok


# ---------------------------------------------------------------------------
# word fallback (python-docx)
# ---------------------------------------------------------------------------


def test_word_read_document_fallback(ws: Path) -> None:
    from docx import Document

    doc_path = ws / "notes.docx"
    doc = Document()
    doc.add_paragraph("Hello from a test document.")
    doc.add_paragraph("Second paragraph.")
    doc.save(str(doc_path))
    result = word_read_document.run(path="notes.docx")
    assert result.ok
    assert "Hello from a test document." in result.output
    assert "Second paragraph." in result.output


def test_word_read_missing_file(ws: Path) -> None:
    result = word_read_document.run(path="nope.docx")
    assert not result.ok and "not found" in result.output


# ---------------------------------------------------------------------------
# office tools are safe when Office is absent (any platform)
# ---------------------------------------------------------------------------


def test_outlook_tools_error_gracefully_without_office(ws: Path) -> None:
    from james.tools.office_tools import (
        outlook_create_event,
        outlook_read_inbox,
        outlook_send_email,
    )

    for result in (
        outlook_read_inbox.run(count=2),
        outlook_send_email.run(to="a@b.c", subject="s", body="b"),
        outlook_create_event.run(subject="meet", start="2027-01-01T10:00", end="2027-01-01T11:00"),
    ):
        assert isinstance(result, object)
        assert hasattr(result, "ok")
        assert result.ok is False, result.output
        assert result.output


def test_powerpoint_tool_builds_or_errors(ws: Path) -> None:
    from james.tools.office_tools import powerpoint_create

    result = powerpoint_create.run(title="Deck")
    # python-pptx is a packaged dependency: it builds a real .pptx. When the
    # package is missing the tool must degrade to a clear error.
    assert result.ok is True or (result.ok is False and "pptx" in result.output.lower())


# ---------------------------------------------------------------------------
# notify tool (plyer may be absent in CI)
# ---------------------------------------------------------------------------


def test_notify_tool_missing_plyer_is_graceful() -> None:
    from james.tools.system_tools import notify

    if "plyer" not in sys.modules:
        result = notify.run(title="T", message="M")
        assert result.ok is False and "plyer" in result.output.lower()


def test_notify_tool_with_fake_plyer(monkeypatch: pytest.MonkeyPatch) -> None:
    from james.tools.system_tools import notify

    sent: list[dict] = []
    fake_plyer = SimpleNamespace(notification=SimpleNamespace(notify=lambda **kw: sent.append(kw)))
    monkeypatch.setitem(sys.modules, "plyer", fake_plyer)
    result = notify.run(title="Hello", message="World")
    assert result.ok
    assert sent and sent[0]["title"] == "Hello" and sent[0]["message"] == "World"


# ---------------------------------------------------------------------------
# cloud plugin registry
# ---------------------------------------------------------------------------


@pytest.fixture
def isolated_marketplace(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> Path:
    import james.tools.marketplace as marketplace_module

    market_file = tmp_path / "marketplace.json"
    monkeypatch.setattr(marketplace_module, "_MARKETPLACE_FILE", market_file)
    monkeypatch.setattr(workspace.settings.assistant, "workspace_dir", tmp_path)
    monkeypatch.setattr(marketplace_module.settings.assistant, "workspace_dir", tmp_path)
    return market_file


def _fake_get(payload):
    def _get(url: str, **kwargs):
        return SimpleNamespace(json=lambda: payload, raise_for_status=lambda: None)

    return _get


def test_sync_remote_catalog_merges(
    isolated_marketplace: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    import james.tools.marketplace as marketplace_module

    remote = [
        {"name": "cloud-plugin-a", "description": "A", "version": "1.0.0"},
        {"name": "cloud-plugin-b", "description": "B", "version": "0.2.0"},
        {"broken": "entry"},
    ]
    monkeypatch.setattr(
        "james.tools.marketplace.requests.get",
        _fake_get(remote),
    )
    result = marketplace_module.sync_remote_catalog(url="https://example.invalid/plugins.json")
    assert result["ok"] is True
    assert result["added"] == 2  # the malformed entry is skipped
    catalog = json.loads(isolated_marketplace.read_text(encoding="utf-8"))
    by_name = {e["name"]: e for e in catalog}
    assert by_name["cloud-plugin-a"]["source"] == "remote"
    assert by_name["cloud-plugin-b"]["source"] == "remote"
    assert result["total"] == len(catalog)


def test_sync_remote_catalog_keeps_local_entries(
    isolated_marketplace: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    import james.tools.marketplace as marketplace_module

    isolated_marketplace.write_text(json.dumps([{"name": "local-plugin", "description": "L"}]))
    monkeypatch.setattr(
        "james.tools.marketplace.requests.get",
        _fake_get([{"name": "remote-x", "description": "R"}]),
    )
    result = marketplace_module.sync_remote_catalog(url="https://example.invalid/plugins.json")
    assert result["ok"]
    names = [e["name"] for e in json.loads(isolated_marketplace.read_text(encoding="utf-8"))]
    assert "local-plugin" in names and "remote-x" in names


def test_sync_remote_catalog_network_failure(
    isolated_marketplace: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    import james.tools.marketplace as marketplace_module

    def _fail(url: str, **kwargs):
        raise ConnectionError("network down")

    monkeypatch.setattr("james.tools.marketplace.requests.get", _fail)
    result = marketplace_module.sync_remote_catalog(url="https://example.invalid/plugins.json")
    assert result["ok"] is False
    assert "network" in result["error"].lower()
    assert not isolated_marketplace.exists()  # local catalog untouched


def test_sync_remote_catalog_bad_shape(
    isolated_marketplace: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    import james.tools.marketplace as marketplace_module

    monkeypatch.setattr("james.tools.marketplace.requests.get", _fake_get({"not": "a list"}))
    result = marketplace_module.sync_remote_catalog(url="x")
    assert result["ok"] is False and "JSON list" in result["error"]


def test_marketplace_status_tracks_last_sync(
    isolated_marketplace: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    import james.tools.marketplace as marketplace_module

    monkeypatch.setattr(
        "james.tools.marketplace.requests.get", _fake_get([{"name": "p", "description": "d"}])
    )
    marketplace_module.sync_remote_catalog(url="https://example.invalid/plugins.json")
    status = marketplace_module.marketplace_status()
    assert status["remote_count"] == 1
    assert status["synced_at"] is not None
    assert status["url"].startswith("http")


def test_update_marketplace_tool(
    isolated_marketplace: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    from james.tools.marketplace import update_marketplace

    monkeypatch.setattr(
        "james.tools.marketplace.requests.get", _fake_get([{"name": "p", "description": "d"}])
    )
    result = update_marketplace.run()
    assert result.ok and "Merged 1" in result.output


def test_update_marketplace_tool_failure(
    isolated_marketplace: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    from james.tools.marketplace import update_marketplace

    def _fail(url: str, **kwargs):
        raise ConnectionError("down")

    monkeypatch.setattr("james.tools.marketplace.requests.get", _fail)
    result = update_marketplace.run()
    assert not result.ok and "down" in result.output


# ---------------------------------------------------------------------------
# registry surface
# ---------------------------------------------------------------------------


def test_marketplace_tools_registered() -> None:
    registry = ToolRegistry(discover_plugins=False)
    assert "update_marketplace" in registry.names()
